from docx import Document

duplicate_tables = []


def delete_tables(table_list):
    sorted_table_list = sorted(table_list, reverse=True)
    print(sorted_table_list)
    for table_index in sorted_table_list:
        document.tables[table_index]._element.getparent().remove(
            document.tables[table_index]._element
        )


# Returns True if the tables are the same
def compare_tables(table1, table2):
    if len(table1.rows) != len(table2.rows):
        return False

    for row_index in range(len(table1.rows)):
        row1 = table1.rows[row_index]
        row2 = table2.rows[row_index]

        if len(row1.cells) != len(row2.cells):
            return False

        for cell_index in range(len(row1.cells)):
            cell1 = row1.cells[cell_index]
            cell2 = row2.cells[cell_index]

            if cell1.text != cell2.text:
                return False

    return True


def delete_duplicate_tables(document):
    for table_index in range(len(document.tables)):
        if table_index in duplicate_tables:
            continue
        for comparison_table_index in range(table_index + 1, len(document.tables)):
            if comparison_table_index in duplicate_tables:
                continue
            if compare_tables(
                document.tables[table_index], document.tables[comparison_table_index]
            ):
                duplicate_tables.append(comparison_table_index)

    delete_tables(duplicate_tables)


original_doc_name = "./original.docx"  # Path to your original docx file
new_doc_name = "./new.docx"  # Path to your new docx file (could be the same if you want the original to be overwritten)
document = Document(original_doc_name)

delete_duplicate_tables(document)

document.save(new_doc_name)
